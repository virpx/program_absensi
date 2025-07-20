import time
import socket
import pyqrcode
import tkinter as tk
from tkinter import ttk, messagebox
from tkinter import font as tkFont
from tkinter import font
from tkinter import messagebox ,filedialog
from PIL import Image, ImageTk
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
import io
import re
import pandas
import requests
from tkinter import filedialog
from openpyxl import load_workbook
from docx import Document
from datetime import datetime
import os
import json
import base64
import pandas as pd
from pathlib import Path
import os
import subprocess
import mysql.connector
from mysql.connector import Error
def pilih_folder(parametertitle):
    root = tk.Tk()
    root.withdraw()  # Menyembunyikan jendela utama
    folder = filedialog.askdirectory(title=parametertitle)  # Membuka dialog untuk memilih folder
    return folder
# Inisialisasi driver di luar fungsi agar dapat diakses secara global
driver = None
# pilih_aksi = tk.StringVar()
btnsave = None
def show_frame(frame,tambahan=""):
    # Sembunyikan semua frame
    if tambahan == "kereport":
        getdatareport(1)
    elif tambahan == "keupdate":
        loaddataedit()
    for child in root.winfo_children():
        if isinstance(child, tk.Frame):
            child.pack_forget()
    # Tampilkan frame yang diminta
    frame.pack(fill="both", expand=True)

list_entry_add_siswa = []
tokenlogin = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJleHAiOjE3MzY2NDcyNDAsImRhdGEiOiJhZG1pbmxvZ2luIiwiaWF0IjoxNzM2NTYwODQwfQ.nUn4QFwha9pt9VScgHY2ekqVc_BK8hsCUOG2q_1_o58"
# Fungsi untuk mengecek login setiap beberapa detik
def check_login_status():
    if is_logged_in():
        show_frame(absen_frame)
    else:
        # Cek lagi dalam 5 detik
        root.after(5000, check_login_status)

# Fungsi yang dipanggil saat jendela Tkinter ditutup
def on_closing():
    global driver
    if driver:
        driver.quit()
    root.destroy() 

# Cek Login WA
def is_logged_in():
    try:
        # Coba cari elemen yang hanya ada jika login berhasil
        chat_list = driver.find_elements("css selector", "#side")
        return len(chat_list) > 0
    except Exception as e:
        return False
# Cek qrcode absen
arrpesankirim = ["","",""]
def kirim_notifikasi_presensi(no_hp,nama,status,keterangan,waktuhadir):
    pesankirim = arrpesankirim[status].replace("{{nama}}","" if nama == None else nama).replace("{{tanggal}}","" if waktuhadir == None else waktuhadir).replace("{{keteranganizin}}","" if keterangan == None else keterangan)
    try:
        print(no_hp+"-"+pesankirim)
        # driver.get(f"https://web.whatsapp.com/send?phone=62{no_hp}&text={pesankirim}")
        driver.execute_script("""
if(document.getElementById("bukachatbaru")){
  document.getElementById("bukachatbaru").setAttribute('href',"https://web.whatsapp.com/send?phone=62"""+str(no_hp)+"""&text="""+str(pesankirim)+"""")
}else{
  linkchat = document.createElement("a")
  linkchat.setAttribute('href',"https://web.whatsapp.com/send?phone=62"""+str(no_hp)+"""&text="""+str(pesankirim)+"""")
  linkchat.setAttribute("id","bukachatbaru")
  document.getElementsByTagName("h1")[0].append(linkchat)
}
document.getElementById("bukachatbaru").click()
        """)
        # Klik tombol kirim pada WhatsApp Web (dengan Xpath)
        while True:
            try:
                send_button = driver.find_element(By.CSS_SELECTOR, 'button[aria-label="Kirim"]').click()
                break
            except Exception as e:
                try:
                    send_button = driver.find_element(By.CSS_SELECTOR, 'button[aria-label="Send"]').click()
                    break
                except Exception as e:
                    pass
                    adaerror = False
                    if "Phone number shared via url is invalid." in driver.find_element(By.TAG_NAME,"html").text:
                        print(driver.find_element(By.TAG_NAME,"html").text)
                        adaerror = True
                    if adaerror :
                        jalane = True
                        while jalane:
                            try:
                                send_button = driver.find_element(By.CSS_SELECTOR, 'button[aria-label="Kirim"]').click()
                                return True
                            except Exception as e:
                                try:
                                    send_button = driver.find_element(By.CSS_SELECTOR, 'button[aria-label="Send"]').click()
                                    return True
                                except Exception as e:
                                    for a in driver.find_elements(By.TAG_NAME,"button"):
                                        try:
                                            if "OK" in a.text:
                                                a.click()
                                                return False
                                        except:
                                            pass
        return True
    except Exception as e:
        messagebox.showerror("Error", f"Gagal membuka WhatsApp Web: {str(e)}")
        return False
def sanitize_phone_number(phone):
    phone = re.sub(r'\D', '', phone)
    if phone.startswith('+62'):
        phone = phone[3:]  
    elif phone.startswith('62'):
        phone = phone[2:]  
    elif phone.startswith('0'):
        phone = phone[1:]  
    return str(phone)
def create_rounded_button(canvas, x, y, width, height, radius, text, command=None):
    canvas.create_oval(x, y, x + radius*2, y + radius*2, fill="purple", outline="")
    canvas.create_oval(x + width - radius*2, y, x + width, y + radius*2, fill="purple", outline="")
    canvas.create_oval(x, y + height - radius*2, x + radius*2, y + height, fill="purple", outline="")
    canvas.create_oval(x + width - radius*2, y + height - radius*2, x + width, y + height, fill="purple", outline="")
    canvas.create_rectangle(x + radius, y, x + width - radius, y + height, fill="purple", outline="")
    canvas.create_rectangle(x, y + radius, x + width, y + height - radius, fill="purple", outline="")
    button_text = canvas.create_text(x + width / 2, y + height / 2, text=text, fill="white", font=("Arial", 12, "bold"))
    if command:
        canvas.tag_bind(button_text, "<Button-1>", lambda event: command())
        canvas.tag_bind("button_background", "<Button-1>", lambda event: command())
# Fungsi untuk memuat gambar dari folder assets
def load_image(path, width, height):
    image = Image.open(path)
    image = image.resize((width, height), Image.LANCZOS)
    return ImageTk.PhotoImage(image)

# Membuat antarmuka GUI menggunakan Tkinter
root = tk.Tk()
root.title("WhatsApp Web QR Code Fetcher")

# Ukuran window
root.geometry("1000x1000")
# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame Login # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #


def login():
    global tokenlogin
    userid = email_entry.get()
    password = password_entry.get()
    payload = {
        "username": userid,
        "password":password
    }
    response = requests.post("http://localhost:3000/login", data=payload)
    datalogin = json.loads(response.text)
    if datalogin["success"] == 0:
        messagebox.showerror("Login Failed", datalogin["data"])
    else:
        # Pindah ke frame QR Code
        tokenlogin = datalogin["data"]
        email_entry.delete(0, tk.END)
        password_entry.delete(0, tk.END)
        show_frame(menu_frame)
        # fetch_and_show_qr_code()

login_frame = tk.Frame(root, bg="white")

# Sub-frame untuk memusatkan isi
center_frame = tk.Frame(login_frame, bg="white")
center_frame.pack(expand=True, padx=20, pady=20)

# Gambar di tengah
try:
    original_image = Image.open(str(Path(__file__).resolve().parent)+"\\assets\\login_image.png")
    resized_image = original_image.resize((250, 250))  # Sesuaikan ukuran gambar
    photo = ImageTk.PhotoImage(resized_image)
    image_label = tk.Label(center_frame, image=photo, borderwidth=0, highlightthickness=0, bg="white")
    image_label.grid(row=0, column=0, padx=20, rowspan=3)  # Gambar menempati posisi kiri
except Exception as e:
    error_label = tk.Label(center_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.grid(row=0, column=0, padx=20, rowspan=3)

# Sub-frame kanan untuk form login
form_frame = tk.Frame(center_frame, bg="white")
form_frame.grid(row=0, column=1, padx=20, sticky="n")

# Judul
label_font = tkFont.Font(family="Helvetica", size=20, weight="bold")
title_label = tk.Label(form_frame, text="Login as Admin", font=label_font, bg="white", fg="#6F42C1")
title_label.pack(pady=(20, 20))

# Frame untuk input email
email_frame = tk.Frame(form_frame, bg="white")
email_frame.pack(pady=(10, 10), fill="x")

# Email icon
email_icon = tk.Label(email_frame, text="👤", font=("Helvetica", 14), bg="white", fg="#6F42C1")
email_icon.pack(side=tk.LEFT, padx=(10, 10))

# Email entry
email_entry = tk.Entry(
    email_frame,
    font=("Helvetica", 14),
    bg="#f0f0ff",
    fg="black",
    insertbackground="black",
    highlightthickness=2,
    highlightbackground="#a0a0ff",
    relief="solid",
    bd=1
)
email_entry.pack(side=tk.LEFT, fill="x", expand=True, padx=(0, 10), ipady=5)

# Frame untuk input password
password_frame = tk.Frame(form_frame, bg="white")
password_frame.pack(pady=(10, 20), fill="x")

# Password icon
password_icon = tk.Label(password_frame, text="🔒", font=("Helvetica", 14), bg="white", fg="#6F42C1")
password_icon.pack(side=tk.LEFT, padx=(10, 10))

# Password entry
password_entry = tk.Entry(
    password_frame,
    font=("Helvetica", 14),
    bg="#f0f0ff",
    fg="black",
    insertbackground="black",
    highlightthickness=2,
    highlightbackground="#a0a0ff",
    relief="solid",
    bd=1,
    show="*"
)
password_entry.pack(side=tk.LEFT, fill="x", expand=True, padx=(0, 10), ipady=5)

# Login button
button_frame = tk.Frame(form_frame, bg="white")  # Sub-frame untuk memisahkan tombol
button_frame.pack(pady=(5, 0), fill="x")  # Tambahkan padding di atas tombol

# Tombol login
login_btn = tk.Canvas(button_frame, width=110, height=50, bg="white", highlightthickness=0)
login_btn.pack(anchor="center", pady=(4, 0))  # Centering tombol login dan beri padding bawah
create_rounded_button(
    login_btn,
    x=5,
    y=5,
    width=100,
    height=40,
    radius=20,
    text="LOGIN",
    command=login
)

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame Menu # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #


# Fungsi untuk menampilkan QR Code dan cek login       
def from_menu_to_qrcode():
    requests.get("http://localhost:3000/generateabsen")
    show_frame(qr_frame)
    fetch_and_show_qr_code()
    check_login_status()

menu_frame = tk.Frame(root, bg="white")

# Palet warna utama
primary_color = "#6A1B9A"  # Ungu untuk aksen utama
background_color = "white"

# Section Tengah (terdiri dari kiri, tengah, dan kanan)
tengah_frame = tk.Frame(menu_frame, bg="white")
tengah_frame.place(relx=0.5, rely=0.4, anchor="center", relwidth=0.9, relheight=0.5)

# Bagian Kiri di Tengah
kiri_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
kiri_frame.pack(fill="y", side="left", padx=10, pady=10)

# Bagian Tengah di Tengah
tengah_tengah_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
tengah_tengah_frame.pack(fill="both", side="left", expand=True, padx=10, pady=10)

# Bagian Kanan di Tengah
kanan_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
kanan_frame.pack(fill="y", side="right", padx=10, pady=10)

# Section Bawah
bawah_frame = tk.Frame(menu_frame, bg=background_color, height=80)
bawah_frame.pack(fill="x", side="bottom")

# Judul Menu
title_label = tk.Label(menu_frame, text="Main Menu", bg=background_color, fg="Black", font=("Arial", 24, "bold"))
title_label.pack(pady=10)

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\database.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(kiri_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(kiri_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Tombol Database di halaman Menu
database = tk.Canvas(kiri_frame, width=110, height=50, bg="white", highlightthickness=0)
database.place(relx=0.50, rely=0.85, anchor="center") 
create_rounded_button(database, x=5, y=5, width=100, height=40, radius=20, text="Database", command=lambda: show_frame(databases_frame))

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\reicon.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(tengah_tengah_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(tengah_tengah_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Tombol Riport di halaman Menu
riport = tk.Canvas(tengah_tengah_frame, width=110, height=50, bg="white", highlightthickness=0)
riport.place(relx=0.50, rely=0.85, anchor="center") 
create_rounded_button(riport, x=5, y=5, width=100, height=40, radius=20, text="Report", command=lambda: show_frame(report_frame,"kereport"))

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\abs.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(kanan_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(kanan_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Tombol Absensi Siswa di halaman Menu
absensi = tk.Canvas(kanan_frame, width=110, height=50, bg="white", highlightthickness=0)
absensi.place(relx=0.50, rely=0.85, anchor="center") 
create_rounded_button(absensi, x=5, y=5, width=100, height=40, radius=20, text="Absensi Siswa", command=from_menu_to_qrcode)

# Tombol Ubah Pesan di pojok kiri bawah pada bawah_frame
ubah_pesan_canvas = tk.Canvas(bawah_frame, width=120, height=50, bg="white", highlightthickness=0)
ubah_pesan_canvas.place(relx=0.15, rely=0.25, anchor="center")  # Kiri bawah, relx lebih kecil dari Exit
create_rounded_button(ubah_pesan_canvas, x=5, y=5, width=110, height=40, radius=20, text="Ubah Pesan", command=lambda: (show_frame(ubahpesan_frame),loadpesanwa()))

# Tombol Exit di pojok kanan bawah pada bawah_frame
exited_canvas = tk.Canvas(bawah_frame, width=100, height=50, bg="white", highlightthickness=0)
exited_canvas.place(relx=0.85, rely=0.25, anchor="center") 
create_rounded_button(exited_canvas, x=5, y=5, width=90, height=40, radius=20, text="Exit", command=lambda: show_frame(login_frame))


# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame QR Code # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #

# Fungsi untuk mengambil dan menampilkan QR code di Tkinter
def fetch_and_show_qr_code():
    global driver, qr_label
    try:
        # Konfigurasi headless browser menggunakan Chrome
        chrome_options = Options()
        # chrome_options.add_argument("--headless")
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument('--disable-dev-shm-usage')

        # Inisialisasi WebDriver
        driver = webdriver.Chrome(options=chrome_options)
        driver.get("https://web.whatsapp.com/")
        
        # Tunggu sebentar agar halaman selesai dimuat
        time.sleep(4)
        
        # Tunggu elemen dengan atribut data-ref muncul
        while True:
            try:
                qr_element = driver.find_element("css selector", "div[data-ref]")
                time.sleep(1)
                break
            except:
                pass

        # Dapatkan nilai dari atribut data-ref
        data_ref = qr_element.get_attribute("data-ref")
        
        # Buat QR code dari data-ref
        qr_code = pyqrcode.create(data_ref)

        # Simpan QR code ke dalam stream gambar PNG
        buffer = io.BytesIO()
        qr_code.png(buffer, scale=6)
        buffer.seek(0)

        # Buka gambar menggunakan PIL dan konversi ke format yang bisa digunakan di Tkinter
        qr_image = Image.open(buffer)
        qr_image_tk = ImageTk.PhotoImage(qr_image)

        # Jika ada QR code yang sudah ditampilkan sebelumnya, hapus
        if qr_label is not None:
            qr_label.destroy()

        # Label untuk menampilkan gambar QR code di Tkinter
        qr_label = tk.Label(qr_frame, image=qr_image_tk)
        qr_label.image = qr_image_tk  
        qr_label.pack(pady=20)

    except Exception as e:
        messagebox.showerror("Error", f"Gagal mengambil QR Code: {str(e)}")


qr_frame = tk.Frame(root , bg="white")

# Variabel untuk label gambar QR code
qr_label = None

# Menambahkan deskripsi
description_label = tk.Label(qr_frame, text="Scan the QR code above to proceed.", 
                              bg="white", fg="black", font=("Arial", 10))
description_label.pack(pady=10)

# Tombol Back di halaman QR Code
back_button = tk.Button(qr_frame, text="Back", command=lambda: show_frame(menu_frame) , bg="#4CAF50", fg="black", font=("Arial", 12, "bold"))
back_button.pack(pady=20)


# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame Absensi # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #
def normalize_phone_number(phone_number):
    # Hilangkan semua karakter kecuali angka dan +
    phone_number = re.sub(r'[^\d+]', '', phone_number)
    
    # Normalisasi awalan
    if phone_number.startswith('+62'):
        phone_number = phone_number[3:]
    elif phone_number.startswith('62'):
        phone_number = phone_number[2:]
    elif phone_number.startswith('0'):
        phone_number = phone_number[1:]
    
    return phone_number

def is_valid_indonesian_number(phone_number):
    normalized = normalize_phone_number(phone_number)
    print(f"Normalized: {normalized}")  # DEBUG
    # Validasi: mulai dari 8, total 9-12 digit
    pattern = r'^8[1-9][0-9]{7,10}$'
    return bool(re.fullmatch(pattern, normalized))

def getabsenhariinikirim():
    global tokenlogin
    payload = {
        "login": tokenlogin,
        "data":"g"
    }
    response = requests.get("http://localhost:3000/getdatanotifikasi", data=payload)
    dataabsenhariini = json.loads(response.text)
    loadpesanwa()
    for a in range(0,len(datapesanwa)):
        arrpesankirim[a] = datapesanwa[a]["isi"]
    for a in dataabsenhariini:
        if is_valid_indonesian_number(str(a["no_ortu"])):
            kirim_notifikasi_presensi(normalize_phone_number(a["no_ortu"]),a["nama"],a["status"],a["keterangan"],a["waktuhadir"])
        else:
            print("Invalid Number Phone")

def on_key_release(e):
    global tokenlogin
    if len(nik_entry.get()) == 10:
        payload = {
            "login": tokenlogin,
            "data":nik_entry.get()
        }
        response = requests.post("http://localhost:3000/absen", data=payload)
        datalogin = json.loads(response.text)
        if datalogin['success'] == 0:
            messagebox.showerror("Gagal Absen", datalogin["data"])
        nik_entry.delete(0,tk.END)

absen_frame = tk.Frame(root,bg="white")

# Membuat dua section di dalam frame absensi

# Membuat frame atas dan bawah untuk menambah margin vertikal pada garis pemisah
top_margin = tk.Frame(absen_frame, height=20, bg="white")  # margin atas
top_margin.pack(fill="x")
section_kiri = tk.Frame(absen_frame, width=550, bg="white")
section_kiri.pack(side="left", fill="both", expand=True)

section_kanan = tk.Frame(absen_frame, width=250, bg="white")
section_kanan.pack(side="right", fill="both", expand=True)

# Membuat garis tegak lurus berwarna ungu sebagai pemisah
separator = tk.Canvas(absen_frame, width=20, height=560, bg="white", highlightthickness=0)
separator.create_line(10, 0, 10, 560, fill="purple", width=5)  # Garis ungu di tengah canvas
separator.pack(side="left", padx=10, pady=20)

# Isi section kiri

# NIK scan
nik = tk.Label(section_kiri, text="Barcode Scan", font=("Arial", 26, "bold"), fg="black", bg="white" )
nik.pack(pady=(100, 10), padx=(10, 0),anchor="center")

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\scan.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 350, 350)  # Mengatur ukuran gambar
    img_label = tk.Label(section_kiri, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(section_kiri, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Isi section kanan

# Entry untuk input nomor
nik_entry = tk.Entry(
    section_kanan,
    font=("Helvetica", 14),       
    bg="#f0f0ff",                  
    fg="black",                    
    insertbackground="black",      
    highlightthickness=2,          
    highlightbackground="#a0a0ff", 
    relief="solid",                
    bd=1                           
)

nik_entry.place(relx=0.5, rely=0.4, anchor="center", width=200, height=30)  # Ukuran dan posisi

# Bind event key release pada entry
nik_entry.bind("<KeyRelease>", on_key_release)

# Membuat Canvas untuk tombol di section kanan
back_canvas = tk.Canvas(section_kanan, width=100, height=50, bg="white", highlightthickness=0)
back_canvas.place(relx=0.5, rely=0.6, anchor="center")  # Posisi di tengah secara horizontal, sedikit di bawah tengah secara vertikal

# Membuat tombol "Back" melingkar
create_rounded_button(back_canvas, x=5, y=5, width=90, height=40, radius=20, text="Back", command=lambda: show_frame(menu_frame))
wa_canvas = tk.Canvas(section_kanan, width=140, height=50, bg="white", highlightthickness=0)
wa_canvas.place(relx=0.5, rely=0.75, anchor="center")  # Posisi di tengah, sedikit lebih bawah dari tombol "Back"

# Membuat tombol "KIRIM WA" melingkar
create_rounded_button(wa_canvas, x=5, y=5, width=130, height=40, radius=20, text="KIRIM WA", command=lambda: getabsenhariinikirim())

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame Database # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #


# Palet warna utama
primary_color = "#6A1B9A"  # Ungu untuk aksen utama
background_color = "white"

# Frame ubah pesan
ubahpesan_frame = tk.Frame(root, bg="white")

# Style tabel
sty = ttk.Style()
sty.theme_use("default")
sty.configure("Treeview",
                background="#6A1B9A",
                foreground="black",
                rowheight=40,
                fieldbackground="white",
                font=('Arial', 12))
sty.configure("Treeview.Heading",
                font=('Arial', 14, 'bold'),
                background="#6A1B9A",
                foreground="white")
sty.map('Treeview', background=[('selected', '#9575CD')])
datapesanwa = None
indexeditpesanwa = -1;
# Text Field Pesan
pesan_var = tk.StringVar()
top_frame = tk.Frame(ubahpesan_frame, bg="white")
top_frame.pack(pady=20)

pesan_entry = tk.Entry(top_frame, textvariable=pesan_var, font=("Arial", 14), width=50, relief="solid", bd=2)
pesan_entry.pack(side="left", padx=(0,10))

def save_pesan():
    pesan = pesan_var.get().strip()
    if not pesan:
        messagebox.showwarning("Peringatan", "Pesan tidak boleh kosong!")
        return
    for a in range(0,len(datapesanwa)):
        if datapesanwa[a]["status"] == indexeditpesanwa:
            datapesanwa[a]["isi"] = pesan
            break
    pesan_var.set("")
    json_str = json.dumps(datapesanwa)
    json_bytes = json_str.encode('utf-8')

    # Step 3: Encode to Base64
    base64_bytes = base64.b64encode(json_bytes)
    base64_str = base64_bytes.decode('utf-8')
    payload = {
        "login": tokenlogin,
        "data":base64_str
    }
    response = requests.post("http://localhost:3000/simpanpesanwa", data=payload)
    if response.status_code == 200:
        clear_tree()
        loadpesanwa()
        clear_tree()
        loadpesanwa()
    else:
        print("Gagal menyimpan pesan. Status code:", response.status_code)
def loadpesanwa():
    global datapesanwa
    payload = {
        "login": tokenlogin,
    }
    hasilloadpesan = requests.get("http://localhost:3000/getpesanwa", data=payload)
    datapesanwa = json.loads(hasilloadpesan.text)["data"]
    for a in datapesanwa:
        insert_row(a["isi"],"Bolos" if a["status"] == 0 else "Masuk" if a["status"] == 1 else "Izin")
def clear_tree():
    for item in pesan_table.get_children():
        pesan_table.delete(item)
def on_double_click(event):
    global indexeditpesanwa
    selected_item = pesan_table.selection()
    if selected_item:
        values = pesan_table.item(selected_item[0], "values")
        indexeditpesanwa = 0 if values[1] == "Bolos" else 1 if values[1] == "Masuk" else 2
        pesan_var.set(values[0])
save_button = tk.Button(top_frame, text="Save", command=save_pesan,
                        bg="#6A1B9A", fg="white", font=("Arial", 12, "bold"),
                        relief="raised", bd=3, padx=15, pady=5)
save_button.pack(side="left")

# Tombol Tambah Nama, Waktu, Tanggal
button_frame = tk.Frame(ubahpesan_frame, bg="white")
button_frame.pack(pady=(0,20))

def tambah_tag(tag):
    current_text = pesan_var.get()
    pesan_var.set(current_text + f" {{{{{tag}}}}}")

for idx, (label, tag) in enumerate([("Nama", "nama"), ("Keterangan Izin", "keteranganizin"), ("Tanggal", "tanggal")]):
    btn = tk.Button(button_frame, text=label, command=lambda t=tag: tambah_tag(t),
                    bg="#6A1B9A", fg="white", font=("Arial", 10, "bold"),
                    relief="raised", bd=2, padx=10, pady=5)
    btn.grid(row=0, column=idx, padx=10)

# Frame tabel + scrollbar
table_frame = tk.Frame(ubahpesan_frame, bg="white")
table_frame.pack(pady=10)

columns = ("Pesan", "Aksi")
pesan_table = ttk.Treeview(table_frame, columns=columns, show="headings", height=8)
pesan_table.heading("Pesan", text="Pesan")
pesan_table.heading("Aksi", text="Aksi")
pesan_table.column("Pesan", width=600, anchor="w")
pesan_table.column("Aksi", width=200, anchor="center")
pesan_table.pack(side="left")
pesan_table.bind("<Double-1>", on_double_click)
# Scrollbar
scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=pesan_table.yview)
pesan_table.configure(yscrollcommand=scrollbar.set)
scrollbar.pack(side="right", fill="y")

# Tempat simpan frame button per-row
action_widgets = {}

# Fungsi Insert Row dengan Button Update & Delete
def insert_row(pesan,status):
    item_id = pesan_table.insert("", "end", values=(pesan,status))

def exit_button_pressed():
    # Menyembunyikan ubahpesan_frame
    ubahpesan_frame.pack_forget()
    
    # Menampilkan menu_frame
    show_frame(menu_frame)

# Section Bawah
bawah_frame = tk.Frame(ubahpesan_frame, bg=background_color, height=80)
bawah_frame.pack(side="bottom", fill="x", padx=20, pady=10)

# Tombol Exit di pojok kanan bawah pada bawah_frame
exited_canvas = tk.Canvas(bawah_frame, width=100, height=50, bg="white", highlightthickness=0)
exited_canvas.place(relx=0.85, rely=0.25, anchor="center") 
create_rounded_button(exited_canvas, x=5, y=5, width=90, height=40, radius=20, text="Exit",command=lambda: show_frame(menu_frame))

# Frame utama database
databases_frame = tk.Frame(root, bg="white")

# Judul Menu
title_label = tk.Label(databases_frame, text="Menu Database", bg=background_color, fg="Black", font=("Arial", 24, "bold"))
title_label.pack(pady=20)

# Section Tengah (terdiri dari kiri, tengah, dan kanan)
tengah_frame = tk.Frame(databases_frame, bg=background_color)
tengah_frame.pack(fill="both", expand=True, padx=20, pady=20)

# Bagian Kiri di Tengah
kiri_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
kiri_frame.pack(side="left", fill="y", padx=20, pady=20)

# Bagian Tengah di Tengah
tengah_tengah_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
tengah_tengah_frame.pack(side="left", fill="both", expand=True, padx=20, pady=20)

# Bagian Kanan di Tengah
kanan_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
kanan_frame.pack(side="right", fill="y", padx=20, pady=20)

# Section Bawah
bawah_frame = tk.Frame(databases_frame, bg=background_color, height=80)
bawah_frame.pack(side="bottom", fill="x", padx=20, pady=10)


# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\database.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(kiri_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(kiri_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Tombol Database di halaman Menu
database = tk.Canvas(kiri_frame, width=110, height=50, bg="white", highlightthickness=0)
database.place(relx=0.50, rely=0.65, anchor="center") 
create_rounded_button(database, x=5, y=5, width=100, height=40, radius=20, text="Input Data", command=lambda: show_frame(input_frame))

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\reicon.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(tengah_tengah_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(tengah_tengah_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Tombol Riport di halaman Menu
riport = tk.Canvas(tengah_tengah_frame, width=150, height=50, bg="white", highlightthickness=0)
riport.place(relx=0.50, rely=0.65, anchor="center") 
create_rounded_button(riport, x=5, y=5, width=140, height=40, radius=20, text="Backup Database", command=lambda: show_frame(backup_frame))

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\abs.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(kanan_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(kanan_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)

# Tombol Absensi Siswa di halaman Menu
absensi = tk.Canvas(kanan_frame, width=150, height=50, bg="white", highlightthickness=0)
absensi.place(relx=0.50, rely=0.65, anchor="center") 
create_rounded_button(absensi, x=5, y=5, width=140, height=40, radius=20, text="Update Presensi", command=lambda: show_frame(update_frame,"keupdate"))

# Tombol Exit di pojok kanan bawah pada bawah_frame
exited_canvas = tk.Canvas(bawah_frame, width=100, height=50, bg="white", highlightthickness=0)
exited_canvas.place(relx=0.85, rely=0.25, anchor="center") 
create_rounded_button(exited_canvas, x=5, y=5, width=90, height=40, radius=20, text="Exit", command=lambda: show_frame(menu_frame))

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame input data list siswa # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #

def from_inputdata_to_database():
    input_frame.pack_forget()
    databases_frame.pack(fill="both", expand=True) 

base64save = None;
def loadexcel():
    global base64save
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
    )
    if not file_path:
        return  # No file selected

    try:
        workbook = load_workbook(file_path)
        sheet = workbook.active  # Load the active sheet

        # Clear the existing table
        for row in tree.get_children():
            tree.delete(row)
        tree["columns"] = []

        # Load the header and rows
        headers = [cell.value for cell in sheet[1] if cell.value is not None]
        tree["columns"] = headers
        for header in headers:
            tree.heading(header, text=header)
            tree.column(header, width=100, anchor=tk.W)

        data = []
        for row in sheet.iter_rows(min_row=2, values_only=True):
            # Skip rows where all values are None
            if not all(value is None for value in row):
                filtered_row = [value if value is not None else "" for value in row]
                filtered_row_dict = dict(zip(headers, filtered_row))
                tree.insert("", "end", values=filtered_row[:len(headers)])
                data.append(filtered_row_dict)

        # Save data to a JSON file
        json_file_path = "data.json"
        with open(json_file_path, "w", encoding="utf-8") as json_file:
            json.dump(data, json_file, indent=4)

        # Encode JSON file content to Base64
        with open(json_file_path, "rb") as json_file:
            encoded_data = base64.b64encode(json_file.read()).decode("utf-8")

        # Print encoded data to the terminal
        print("Base64 Encoded Data:")
        print(encoded_data)
        base64save = encoded_data

    except Exception as e:
        messagebox.showerror("Error", f"Failed to load file: {e}")
def savedata():
    payload = {
        "login": tokenlogin,
        "data":"get_backup_data"
    }
    response = requests.post("http://localhost:3000/backupkelas9", data=payload)
    response = json.loads(response.text)
    if response["success"] == 1:
        df = pd.DataFrame(response["data"]["list"])
        output_file = "Backup Kelas 9 "+str(response["data"]["tahunajar"])+".xlsx"
        df.to_excel(pilih_folder("Tempat Simpan File Backup Kelas 9")+"/"+output_file, index=False)
        payload = {
            "login": tokenlogin,
            "data":"go_insert_siswa#"+base64save
        }
        response = requests.post("http://localhost:3000/insertlistsiswa", data=payload)
        response = json.loads(response.text)
        if response["success"] == 1:
            messagebox.showinfo("Success", "Berhasil Menambahkan Data")
        else:
            if response["data"] == "err_tahunajar":
                tanyabuat = messagebox.askyesno("Konfirmasi", "Apakah ingin membuat tahun ajar baru?")
                if tanyabuat:
                    payload = {
                        "login": tokenlogin,
                        "data":"tambahtahunajar#"+base64save
                    }
                    response = requests.post("http://localhost:3000/insertlistsiswa", data=payload)
                    response = json.loads(response.text)
                    if response["success"] == 1:
                        messagebox.showinfo("Success", "Berhasil Menambahkan Data")

                        tree.delete(*tree.get_children())
                        show_frame(databases_frame)
                    else:
                        messagebox.showwarning("Error", response["data"])
            elif response["data"] == "err_belum_saatnya":
                tanyabuat = messagebox.askyesno("Konfirmasi", "Belum saatnya melakukan insert data siswa, ingin tetap lanjutkan?")
                if tanyabuat:
                    payload = {
                        "login": tokenlogin,
                        "data":"goinsertsiswaforce#"+base64save
                    }
                    response = requests.post("http://localhost:3000/insertlistsiswa", data=payload)
                    response = json.loads(response.text)
                    if response["success"] == 1:
                        messagebox.showinfo("Success", "Berhasil Menambahkan Data")
                        tree.delete(*tree.get_children())
                        show_frame(databases_frame)
                    else:
                        messagebox.showwarning("Error", response["data"])
    else:
        messagebox.showwarning("Error", response["data"])

# Frame utama
input_frame = tk.Frame(root, bg="white")

# Label Judul
title_label = tk.Label(input_frame, text="Input Data dari Excel", font=("Arial", 16, "bold"), bg="white", fg="#800080")
title_label.pack(pady=10)

# Frame untuk tombol Exit di pojok kanan atas
exit_frame = tk.Frame(input_frame, bg="white")
exit_frame.pack(fill="x", padx=10)
exit_frame.pack_propagate(False)
exit_frame.config(height=50)

# Tombol Exit di pojok kanan atas
exited_canvas = tk.Canvas(exit_frame, width=100, height=40, bg="white", highlightthickness=0)
exited_canvas.pack(side="right", padx=10)
create_rounded_button(exited_canvas, x=5, y=5, width=90, height=30, radius=15, text="Back", command=lambda: show_frame(databases_frame))

# Frame untuk tombol Load dan Save
button_frame = tk.Frame(input_frame, bg="white")
button_frame.pack(pady=10)

# Tombol untuk memuat file Excel dan Save sejajar
load_button = ttk.Button(button_frame, text="Load Excel File", command=loadexcel)
load_button.pack(side="left", padx=5)

save_button = ttk.Button(button_frame, text="Save Data", command=savedata)
save_button.pack(side="left", padx=5)

# Tabel untuk menampilkan data
tree_frame = tk.Frame(input_frame, bg="white")
tree_frame.pack(fill="both", expand=True, pady=10)

tree = ttk.Treeview(tree_frame, show="headings")
tree.pack(side="left", fill="both", expand=True)

# Scrollbar untuk tabel
scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
scrollbar.pack(side="right", fill="y")
tree.configure(yscrollcommand=scrollbar.set)



# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame backup database # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #


# Palet warna utama
primary_color = "#6A1B9A"  # Ungu untuk aksen utama
background_color = "white"

# Frame utama database
backup_frame = tk.Frame(root, bg="white")

# Judul Menu
title_label = tk.Label(backup_frame, text="Back Up Database", bg=background_color, fg="Black", font=("Arial", 24, "bold"))
title_label.pack(pady=20)

# Section Tengah (terdiri dari kiri, tengah, dan kanan)
tengah_frame = tk.Frame(backup_frame, bg=background_color)
tengah_frame.pack(fill="both", expand=True, padx=20, pady=20)

# Bagian Kiri di Tengah
kiri_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
kiri_frame.pack(side="left", fill="y", padx=20, pady=20)

# Bagian Tengah di Tengah
tengah_tengah_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
tengah_tengah_frame.pack(side="left", fill="both", expand=True, padx=20, pady=20)

# Bagian Kanan di Tengah
kanan_frame = tk.Frame(tengah_frame, bg=background_color, width=200)
kanan_frame.pack(side="right", fill="y", padx=20, pady=20)

# Section Bawah
bawah_frame = tk.Frame(backup_frame, bg=background_color, height=80)
bawah_frame.pack(side="bottom", fill="x", padx=20, pady=10)

# Load Gambar
try:
    image_path = str(Path(__file__).resolve().parent)+"\\assets\\datta.png"  # Pastikan jalur sesuai dengan lokasi file gambar Anda
    loaded_image = load_image(image_path, 250, 250)  # Mengatur ukuran gambar
    img_label = tk.Label(tengah_tengah_frame, image=loaded_image, bg="white")
    img_label.image = loaded_image  # Menyimpan referensi agar tidak dihapus oleh garbage collector
    img_label.pack(pady=20)
except FileNotFoundError:
    error_label = tk.Label(tengah_tengah_frame, text="Gambar tidak ditemukan!", fg="red", bg="white")
    error_label.pack(pady=20)
def gobackupdata():
    payload = {
        "login": tokenlogin,
        "data":"get_backup_data"
    }
    response = requests.get("http://localhost:3000/backupdatabase", data=payload)
    response = json.loads(response.text)
    if response["success"] == 1:
        current_date = datetime.now().strftime("%Y-%m-%d")
        df = pd.DataFrame(response["data"])
        output_file = "Backup Database "+current_date+".xlsx"
        df.to_excel(pilih_folder("Pilih Folder Backup Database")+"/"+output_file, index=False)
        messagebox.showinfo("Success","Berhasil Backup Data!")
    else:
        messagebox.showerror("Error",response["data"])
# Tombol Riport di halaman Menu
riport = tk.Canvas(tengah_tengah_frame, width=150, height=50, bg="white", highlightthickness=0)
riport.place(relx=0.50, rely=0.65, anchor="center") 
create_rounded_button(riport, x=5, y=5, width=140, height=40, radius=20, text="Backup Database", command=lambda:gobackupdata())


# Tombol Exit di pojok kanan bawah pada bawah_frame
exited_canvas = tk.Canvas(bawah_frame, width=100, height=50, bg="white", highlightthickness=0)
exited_canvas.place(relx=0.85, rely=0.25, anchor="center") 
create_rounded_button(exited_canvas, x=5, y=5, width=90, height=40, radius=20, text="Exit", command=lambda: show_frame(databases_frame))

# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame update presensi kehadiran siswa # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #

# Data contoh siswa
siswa_data = []
def loaddataedit():
    global siswa_data
    payload = {
        "login": tokenlogin,
        "data":"g"
    }
    response = requests.get("http://localhost:3000/listabsenhariini", data=payload)
    print(response.text)
    datalogin = json.loads(response.text)
    siswa_data = datalogin
    for widget in result_frame.winfo_children():
            widget.destroy()
    for siswa in siswa_data:
        display_siswa(siswa)



# Fungsi untuk mencari siswa
def search_siswa():
    search_query = search_entry.get().lower()
    for widget in result_frame.winfo_children():
        widget.destroy()

    found_siswa = [siswa for siswa in siswa_data if search_query in siswa["nama"].lower()]
    if found_siswa:
        for siswa in found_siswa:
            display_siswa(siswa)
    else:
        tk.Label(result_frame, text="Siswa tidak ditemukan.", fg="red", bg="white").pack(pady=5)
# def gantiabsen():
    
# Fungsi untuk menampilkan siswa dengan checkbox
def display_siswa(siswa):
    siswa_frame = tk.Frame(result_frame, bg="white", pady=10)
    siswa_frame.pack(fill="x", padx=15)

    # Frame untuk nama
    name_frame = tk.Frame(siswa_frame, bg="white")
    name_frame.pack(side="left")
    
    name_label = tk.Label(name_frame, text=siswa["nama"], bg="white", fg="black", 
                         font=("Helvetica", 11, "bold"), width=20, anchor="w")
    name_label.pack(side="left")
    alasanizin = tk.StringVar()
    alasanizin.set(siswa["keterangan"])
    # Frame untuk checkbox
    checkbox_frame = tk.Frame(siswa_frame, bg="white")
    checkbox_frame.pack(side="left", padx=10)

    masuk_var = tk.BooleanVar()
    izin_var = tk.BooleanVar()
    absen_var = tk.BooleanVar()
    
    if siswa['status'] == 0:
        absen_var.set(True)
    elif siswa['status'] == 1:
        masuk_var.set(True)
    elif siswa['status'] == 2:
        izin_var.set(True)

    def on_check(var):
        if var == "masuk":
            izin_var.set(False)
            absen_var.set(False)
            if masuk_var.get() == False:
                masuk_var.set(True)
            keterangan_frame.pack_forget()
        elif var == "izin":
            masuk_var.set(False)
            absen_var.set(False)
            if izin_var.get():
                keterangan_frame.pack(side="left", padx=(10,0))
            else:
                izin_var.set(True)
        elif var == "absen":
            masuk_var.set(False)
            izin_var.set(False)
            keterangan_frame.pack_forget()
            if absen_var.get() == False:
                absen_var.set(True)

    # Style untuk checkbox
    style = ttk.Style()
    style.configure("Custom.TCheckbutton", background="white", font=("Helvetica", 10))

    masuk_checkbox = ttk.Checkbutton(checkbox_frame, text="Masuk", variable=masuk_var, 
                                    command=lambda: on_check("masuk"), style="Custom.TCheckbutton")
    izin_checkbox = ttk.Checkbutton(checkbox_frame, text="Izin", variable=izin_var, 
                                   command=lambda: on_check("izin"), style="Custom.TCheckbutton")
    absen_checkbox = ttk.Checkbutton(checkbox_frame, text="Absen", variable=absen_var, 
                                    command=lambda: on_check("absen"), style="Custom.TCheckbutton")

    masuk_checkbox.pack(side="left", padx=5)
    izin_checkbox.pack(side="left", padx=5)
    absen_checkbox.pack(side="left", padx=5)

    # Frame untuk keterangan
    keterangan_frame = tk.Frame(siswa_frame, bg="white")
    if siswa['status'] == 2:
        keterangan_frame.pack(side="left", padx=(10,0))
    
    keterangan_label = tk.Label(keterangan_frame, text="Keterangan:", bg="white", font=("Helvetica", 10))
    keterangan_label.pack(side="left")
    
    keterangan_izin = tk.Entry(keterangan_frame, font=("Helvetica", 10), width=20, textvariable=alasanizin)
    keterangan_izin.pack(side="left", padx=(5,10))

    def save_attendance():
        datakirim = ""
        if masuk_var.get():
            datakirim = f"{siswa['nisn']}#1"
        elif izin_var.get():
            datakirim = f"{siswa['nisn']}#2#{keterangan_izin.get()}"
        elif absen_var.get():
            datakirim = f"{siswa['nisn']}#0"
            
        payload = {
            "login": tokenlogin,
            "data": datakirim
        }
        response = requests.post("http://localhost:3000/ubahabsen", data=payload)
        datalogin = json.loads(response.text)
        messagebox.showinfo("Status Ubah Presensi", f"Nama: {siswa['nama']}\nStatus: {datalogin['data']}")

    # Frame untuk button save
    button_frame = tk.Frame(siswa_frame, bg="white")
    button_frame.pack(side="right", padx=15)

    style.configure("Save.TButton", 
                   font=("Helvetica", 10),
                   padding=5,
                   background="#4CAF50")
                   
    save_button = ttk.Button(button_frame, text="Save", 
                            command=save_attendance, 
                            style="Save.TButton")
    save_button.pack(side="right")


# Styling tombol dengan ttk
style = ttk.Style()
style.configure("Purple.TButton", background="#800080", foreground="white", font=("Arial", 10))
style.map("Purple.TButton",
          background=[("active", "#5a005a")],
          foreground=[("active", "white")])

# Frame Utama
update_frame = tk.Frame(root, bg="white")

# Frame untuk header yang berisi judul dan tombol exit
header_frame = tk.Frame(update_frame, bg="white")
header_frame.pack(fill="x", pady=10)

# Label Judul
title_label = tk.Label(header_frame, text="Update Kehadiran Siswa", font=("Arial", 16, "bold"), bg="white", fg="#800080")
title_label.pack(side="left", padx=10)

# Tombol Exit di samping title_label
exited_canvas = tk.Canvas(header_frame, width=100, height=40, bg="white", highlightthickness=0)
exited_canvas.pack(side="left", padx=10)
create_rounded_button(exited_canvas, x=5, y=5, width=90, height=30, radius=15, text="Back", command=lambda: show_frame(databases_frame))

# Search Bar
search_frame = tk.Frame(update_frame, bg="white")
search_frame.pack(fill="x", pady=10)

search_label = tk.Label(search_frame, text="Cari Siswa:", bg="white", fg="black")
search_label.pack(side="left", padx=(0, 10))

search_entry = tk.Entry(search_frame)
search_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
stylee = ttk.Style()
stylee.configure("Cari.TButton", foreground="black")
search_button = ttk.Button(search_frame, text="Search", command=search_siswa, style="Cari.TButton")
search_button.pack(side="left")

# Frame untuk menampilkan hasil
result_frame = tk.Frame(update_frame, bg="white")
result_frame.pack(fill="x", pady=10)


# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame Riport # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #

data_report = []
replacements = {
    "<<tanggal_expxort>>": "",
    "<<nama_siswa>>": "",
    "<<kelas_siswa>>": "",
    "<<jumlah_masuk>>": "",
    "<<jumlah_alpha>>": "",
    "<<jumlah_izin>>": "",
    "<<jenis_export>>": "",
}

def replace_text_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, str(value))
def add_table(document, data):
    table = document.add_table(rows=1, cols=len(data[0]))  # Header dari data
    # Tambah header
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(data[0]):
        hdr_cells[idx].text = header
    
    # Tambah data baris
    for row in data[1:]:
        row_cells = table.add_row().cells
        for idx, item in enumerate(row):
            row_cells[idx].text = str(item)
def getdatareport(parameter):
    if parameter == 1:
        replacements["<<jenis_export>>"] = "Mingguan"
    elif parameter == 2:
        replacements["<<jenis_export>>"] = "Bulanan"
    else:
        replacements["<<jenis_export>>"] = "Semester"
    global data_report
    payload = {
        "login": tokenlogin,
        "data":parameter
    }
    response = requests.get("http://localhost:3000/laporan", data=payload)
    data_report = json.loads(response.text)
    table_export_attendance.delete(*table_export_attendance.get_children())
    cntr = 1;
    for row in data_report:
        table_export_attendance.insert("", "end", values=(cntr,row['nama'],row['hadir'],row['izin'],row['alpha']))
        cntr += 1
def export_data():
    tanggal_sekarang = datetime.now()
    tanggal_terformat = tanggal_sekarang.strftime("%d %B %Y")
    bulan_indonesia = {
        "January": "Januari", "February": "Februari", "March": "Maret",
        "April": "April", "May": "Mei", "June": "Juni",
        "July": "Juli", "August": "Agustus", "September": "September",
        "October": "Oktober", "November": "November", "December": "Desember"
    }
    for english, indonesia in bulan_indonesia.items():
        tanggal_terformat = tanggal_terformat.replace(english, indonesia)
    replacements["<<tanggal_expxort>>"] = tanggal_terformat
    folderuntuksave = pilih_folder("Pilih Folder Untuk Menyimpan Report")
    for a in data_report:
        doc = Document(str(Path(__file__).resolve().parent)+"\\assets\\template_report_siswa.docx")
        replacements["<<nama_siswa>>"] = a["nama"]
        replacements["<<kelas_siswa>>"] = a["kelas"]
        replacements["<<jumlah_alpha>>"] = a["alpha"]
        replacements["<<jumlah_izin>>"] = a["izin"]
        replacements["<<jumlah_masuk>>"] = a["hadir"]
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        table_data = [
            ["Tanggal", "Status","Keterangan"],  # Header
        ]
        for b in a["data"]:
            if b["status"] == 1:
                statustulis = "Masuk"
                keterangantulis = "-"
            elif b["status"] == 2:
                statustulis = "Izin"
                keterangantulis = b["keterangan"]
            elif b["status"] == 0:
                statustulis = "Alpha"
                keterangantulis = "-"
            table_data.append([b["untuktanggal"],statustulis,keterangantulis])
        add_table(doc, table_data)
        current_date = datetime.now().strftime("%Y-%m-%d")
        os.makedirs(folderuntuksave+"/export-"+current_date, exist_ok=True)
        doc.save(folderuntuksave+"/export-"+current_date+"/"+replacements["<<nama_siswa>>"]+".docx")
    messagebox.showinfo("Success", "Berhasil Menyimpan Report")

def show_student_detail(event):
    global nisnsiswa
    selected_item = table_export_attendance.selection()[0]
    values = table_export_attendance.item(selected_item)['values']
    student_name = values[1]
    
    for student in data_report:
        if student['nama'] == student_name:
            getdetailabsenanak(student['nisn'])
            break
    show_frame(detail_frame)
    
# Frame utama backup
report_frame = tk.Frame(root, bg="white")

# Frame untuk header dan back button
header_frame = tk.Frame(report_frame, bg="white")
header_frame.pack(fill="x", padx=20, pady=(20,10))

# Back button
back_canvas = tk.Canvas(header_frame, width=100, height=40, bg="white", highlightthickness=0)
back_canvas.pack(side="left")
create_rounded_button(back_canvas, 0, 0, 100, 40, 20, "Back", command=lambda: show_frame(databases_frame))

# Header label
header_label = tk.Label(header_frame, text="Export attendance", font=("Helvetica", 18, "bold"), bg="white", fg="#333")
header_label.pack(side="left", padx=20)

# Time interval label
interval_label = tk.Label(report_frame, text="Jangka Waktu", font=("Helvetica", 12), bg="white")
interval_label.pack(anchor="w", padx=20)

# Frame for time interval buttons
interval_frame = tk.Frame(report_frame, bg="white")
interval_frame.pack(pady=10, padx=20, anchor="w")

btn = tk.Button(interval_frame, text="Mingguan (Default)", font=("Helvetica", 10), bg="#fff", fg="#333", relief="solid", borderwidth=1, padx=10, pady=5,command=lambda:getdatareport(1))
btn.pack(side="left", padx=5)
btn = tk.Button(interval_frame, text="Bulanan", font=("Helvetica", 10), bg="#fff", fg="#333", relief="solid", borderwidth=1, padx=10, pady=5,command=lambda:getdatareport(2))
btn.pack(side="left", padx=5)
btn = tk.Button(interval_frame, text="Semester", font=("Helvetica", 10), bg="#fff", fg="#333", relief="solid", borderwidth=1, padx=10, pady=5,command=lambda:getdatareport(3))
btn.pack(side="left", padx=5)

# Table (Treeview) for attendance data
table_frame = tk.Frame(report_frame, bg="#fef7f5")
table_frame.pack(pady=20, padx=20, fill="both", expand=True)

# Treeview widget
columns = ("No", "Nama", "Hadir", "Ijin", "Alpha")
table_export_attendance = ttk.Treeview(table_frame, columns=columns, show="headings", height=5)
table_export_attendance.pack(fill="both", expand=True)

# Define column headings and widths
for col in columns:
    table_export_attendance.heading(col, text=col)
    table_export_attendance.column(col, anchor="center", width=100)
table_export_attendance.bind('<Double-1>', lambda event: show_student_detail(event))
# Insert sample data


# Tombol Export to PDF
try:
    icon = tk.PhotoImage(file=str(Path(__file__).resolve().parent)+"\\assets\\pdf_icon.png")  # Pastikan file ada
    export_btn = tk.Button(
        report_frame,
        text="Export to PDF",
        image=icon,
        compound="left",
        font=("Helvetica", 12, "bold"),
        borderwidth=1,
        bg="#e0aaff",
        fg="black",
        padx=10,
        pady=5,
        command=lambda: export_data(),
    )
    export_btn.pack(pady=20)
    export_btn.image = icon  # Simpan referensi agar tidak garbage-collected
except Exception as e:
    export_btn = tk.Button(
        report_frame,
        text="Export to PDF",
        font=("Helvetica", 12, "bold"),
        borderwidth=1,
        bg="#e0aaff",
        fg="black",
        padx=10,
        pady=5,
        command=lambda: export_data(),
    )
    export_btn.pack(pady=20)



# # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # Frame Detail Siswa # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # # #

datadetailabsen = []

def getdetailabsenanak(nisn):
    global datadetailabsen
    payload = {
        "login": tokenlogin,
        "data":nisn
    }
    response = requests.get("http://localhost:3000/detailabsen", data=payload)
    datadetailabsen = json.loads(response.text)
    table_detail_absensi.delete(*table_detail_absensi.get_children())
    name_label.config(text="Nama: "+datadetailabsen["data"][0]['nama'])
    class_label.config(text="Kelas: "+datadetailabsen["data"][0]['kelas'])
    for row in datadetailabsen["data"][1]:
        table_detail_absensi.insert("", "end", values=(row['untuktanggal'],row['status'],row['keterangan']))

detail_frame = tk.Frame(root, bg="#f8f9fa", highlightbackground="#ced4da", highlightthickness=1)

title_label = tk.Label(detail_frame, text="Detail Kehadiran", font=("Helvetica", 18, "bold"), bg="white")
title_label.pack(pady=10)

subtitle_frame = tk.Frame(detail_frame, bg="white")
subtitle_frame.pack(anchor="w", padx=20)

name_label = tk.Label(subtitle_frame, text="Nama: Yoni", font=("Arial", 9), bg="white")
name_label.grid(row=0, column=0, sticky="w")
class_label = tk.Label(subtitle_frame, text="Kelas: 9", font=("Arial", 9), bg="white")
class_label.grid(row=1, column=0, sticky="w")

frame_table_detail = tk.Frame(detail_frame, padx=20, pady=10, bg="white")
frame_table_detail.pack(pady=10)

scrollbar = tk.Scrollbar(frame_table_detail)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

table_detail_absensi = ttk.Treeview(frame_table_detail, yscrollcommand=scrollbar.set, 
                                   selectmode="none", columns=("date", "presence", "note"), show="headings")
table_detail_absensi.pack()

scrollbar.config(command=table_detail_absensi.yview)

table_detail_absensi.column("date", anchor=tk.CENTER, width=150)
table_detail_absensi.column("presence", anchor=tk.CENTER, width=100)
table_detail_absensi.column("note", anchor=tk.CENTER, width=200)

table_detail_absensi.heading("date", text="Date")
table_detail_absensi.heading("presence", text="Presence")
table_detail_absensi.heading("note", text="Note")

style = ttk.Style()
style.configure("Treeview.Heading", font=("Helvetica", 12, "bold"))
style.configure("Treeview", rowheight=30, font=("Helvetica", 12))

# Add back button to detail frame
back_button = tk.Button(detail_frame, text="Back", command=lambda: show_frame(report_frame))
back_button.pack(pady=10)



# Menangani event ketika aplikasi ditutup
root.protocol("WM_DELETE_WINDOW", on_closing)
def check_and_create_database(host, user, password, database_name, sql_file_path):
    try:
        # Connect to MySQL server
        connection = mysql.connector.connect(
            host=host,
            user=user,
            password=password
        )

        if connection.is_connected():
            cursor = connection.cursor()
            cursor.execute("SHOW DATABASES;")
            databases = [db[0] for db in cursor.fetchall()]

            if database_name in databases:
                print(f"Database '{database_name}' already exists.")
            else:
                print(f"Database '{database_name}' does not exist. Creating it...")
                # Execute SQL file to create the database
                with open(sql_file_path, 'r') as sql_file:
                    sql_commands = sql_file.read()
                    for command in sql_commands.split(';'):
                        if command.strip():
                            cursor.execute(command)
                print(f"Database '{database_name}' created successfully.")

                # Switch to the target database
                connection.database = database_name

                # Execute the additional SQL command
                insert_sql = "INSERT INTO `tahunajar`(`tahun`) VALUES (2000);"
                cursor.execute(insert_sql)
                connection.commit()
                print("Additional SQL command executed successfully.")

                cursor.close()

    except Error as e:
        print(f"Error: {e}")

    finally:
        if connection.is_connected():
            connection.close()
def is_port_open(port):
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        try:
            s.connect(("127.0.0.1", port))
            return True
        except ConnectionRefusedError:
            return False

def start_mysql_if_needed():
    if not is_port_open(MYSQL_PORT):
        print("MySQL belum berjalan. Menjalankan...")
        subprocess.Popen(
            [MYSQL_BIN_PATH, f"--defaults-file={MYSQL_CONFIG_PATH}"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        time.sleep(3)
        if is_port_open(MYSQL_PORT):
            print("MySQL berhasil dijalankan.")
        else:
            messagebox.showerror("Error", "Gagal menjalankan MySQL.")
            exit(1)
    else:
        print("MySQL sudah berjalan.")

def start_backend_if_needed():
    if not is_port_open(BACKEND_PORT):
        print("Backend belum berjalan. Menjalankan backend.exe...")
        subprocess.Popen(
            [BACKEND_EXE_PATH],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        time.sleep(2)
        if is_port_open(BACKEND_PORT):
            print("Backend berhasil dijalankan.")
        else:
            messagebox.showerror("Error", "Gagal menjalankan backend.")
            exit(1)
    else:
        print("Backend sudah berjalan.")


if __name__ == "__main__":
    # Database configuration
    HOST = "localhost"  # Adjust as needed
    USER = "root"       # Adjust as needed
    PASSWORD = ""       # Adjust as needed
    DATABASE_NAME = "absensi_siswa"  # Replace with your database name
    SQL_FILE_PATH = str(Path(__file__).resolve().parent)+"\\assets\\db.sql"  # Replace with the path to your SQL file
    # Port
    MYSQL_PORT = 3306
    BACKEND_PORT = 3000

    # Path
    MYSQL_BIN_PATH = str(Path(__file__).resolve().parent)+"\\mysql\\bin\\mysqld.exe"
    MYSQL_CONFIG_PATH = str(Path(__file__).resolve().parent)+"\\mysql\\my.ini"
    BACKEND_EXE_PATH = str(Path(__file__).resolve().parent)+"\\backend.exe"
    start_mysql_if_needed()
    check_and_create_database(HOST, USER, PASSWORD, DATABASE_NAME, SQL_FILE_PATH)
    start_backend_if_needed()
    show_frame(login_frame) 

    # # Jalankan aplikasi
    root.mainloop()
