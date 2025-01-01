import tkinter as tk
from tkinter import ttk
import requests
import json
from docx import Document
from datetime import datetime
import os
# Initialize the main application window
root = tk.Tk()
root.title("Export Attendance")
root.geometry("600x400")
root.configure(bg="white")
tokenlogin = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJleHAiOjE3MzU3OTA3ODQsImRhdGEiOiJhZG1pbmxvZ2luIiwiaWF0IjoxNzM1NzA0Mzg0fQ.kNQq4iiTctBgdmWmfkx6ujGQIe7k0wzNSip-9qvEN8c"
data_report = []
replacements = {
    "<<tanggal_expxort>>": "",
    "<<nama_siswa>>": "",
    "<<kelas_siswa>>": "",
    "<<jumlah_masuk>>": "",
    "<<jumlah_alpha>>": "",
    "<<jumlah_izin>>": "",
    "<<jenis_export>>": "",
}
def replace_text_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, str(value))
def add_table(document, data):
    table = document.add_table(rows=1, cols=len(data[0]))  # Header dari data
    # Tambah header
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(data[0]):
        hdr_cells[idx].text = header
    
    # Tambah data baris
    for row in data[1:]:
        row_cells = table.add_row().cells
        for idx, item in enumerate(row):
            row_cells[idx].text = str(item)
def getdatareport(parameter):
    if parameter == 1:
        replacements["<<jenis_export>>"] = "Mingguan"
    elif parameter == 2:
        replacements["<<jenis_export>>"] = "Bulanan"
    else:
        replacements["<<jenis_export>>"] = "Semester"
    global data_report
    payload = {
        "login": tokenlogin,
        "data":parameter
    }
    response = requests.get("http://localhost:3000/laporan", data=payload)
    data_report = json.loads(response.text)
    table_export_attendance.delete(*table_export_attendance.get_children())
    cntr = 1;
    for row in data_report:
        table_export_attendance.insert("", "end", values=(cntr,row['nama'],row['hadir'],row['izin'],row['alpha']))
        cntr += 1
def export_data():
    tanggal_sekarang = datetime.now()
    tanggal_terformat = tanggal_sekarang.strftime("%d %B %Y")
    bulan_indonesia = {
        "January": "Januari", "February": "Februari", "March": "Maret",
        "April": "April", "May": "Mei", "June": "Juni",
        "July": "Juli", "August": "Agustus", "September": "September",
        "October": "Oktober", "November": "November", "December": "Desember"
    }
    for english, indonesia in bulan_indonesia.items():
        tanggal_terformat = tanggal_terformat.replace(english, indonesia)
    replacements["<<tanggal_expxort>>"] = tanggal_terformat
    for a in data_report:
        doc = Document("template_report_siswa.docx")
        replacements["<<nama_siswa>>"] = a["nama"]
        replacements["<<kelas_siswa>>"] = a["kelas"]
        replacements["<<jumlah_alpha>>"] = a["alpha"]
        replacements["<<jumlah_izin>>"] = a["izin"]
        replacements["<<jumlah_masuk>>"] = a["hadir"]
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        table_data = [
            ["Tanggal", "Status","Keterangan"],  # Header
        ]
        for b in a["data"]:
            if b["status"] == 1:
                statustulis = "Masuk"
                keterangantulis = "-"
            elif b["status"] == 2:
                statustulis = "Izin"
                keterangantulis = b["keterangan"]
            elif b["status"] == 0:
                statustulis = "Alpha"
                keterangantulis = "-"
            table_data.append([b["untuktanggal"],statustulis,keterangantulis])
        add_table(doc, table_data)
        current_date = datetime.now().strftime("%Y-%m-%d")
        os.makedirs("hasil_export/"+current_date, exist_ok=True)
        doc.save("hasil_export/"+current_date+"/"+replacements["<<nama_siswa>>"]+".docx")
# Header label
header_label = tk.Label(root, text="Export attendance", font=("Helvetica", 18, "bold"), bg="white", fg="#333")
header_label.pack(pady=(20, 10))

# Time interval label
interval_label = tk.Label(root, text="Jangka Waktu", font=("Helvetica", 12), bg="white")
interval_label.pack(anchor="w", padx=20)

# Frame for time interval buttons
interval_frame = tk.Frame(root, bg="white")
interval_frame.pack(pady=10, padx=20, anchor="w")

btn = tk.Button(interval_frame, text="Mingguan (Default)", font=("Helvetica", 10), bg="#fff", fg="#333", relief="solid", borderwidth=1, padx=10, pady=5,command=lambda:getdatareport(1))
btn.pack(side="left", padx=5)
btn = tk.Button(interval_frame, text="Bulanan", font=("Helvetica", 10), bg="#fff", fg="#333", relief="solid", borderwidth=1, padx=10, pady=5,command=lambda:getdatareport(2))
btn.pack(side="left", padx=5)
btn = tk.Button(interval_frame, text="Semester", font=("Helvetica", 10), bg="#fff", fg="#333", relief="solid", borderwidth=1, padx=10, pady=5,command=lambda:getdatareport(3))
btn.pack(side="left", padx=5)

# Table (Treeview) for attendance data
table_frame = tk.Frame(root, bg="#fef7f5")
table_frame.pack(pady=20, padx=20, fill="both", expand=True)

# Treeview widget
columns = ("No", "Nama", "Hadir", "Ijin", "Alpha")
table_export_attendance = ttk.Treeview(table_frame, columns=columns, show="headings", height=5)
table_export_attendance.pack(fill="both", expand=True)

# Define column headings and widths
for col in columns:
    table_export_attendance.heading(col, text=col)
    table_export_attendance.column(col, anchor="center", width=100)

# Insert sample data

# Export to PDF button
icon = tk.PhotoImage(file="C:/Users/ferp/Desktop/Semester 7/KP/utama/assets/pdf_icon.png")
export_btn = tk.Button(
    root, text="Export to PDF", 
    image=icon, compound="left",
    font=("Helvetica", 12, "bold"),
    borderwidth=1,bg="#f3e5e5", 
    fg="black", padx=10, 
    pady=5, 
    command=lambda:export_data())
export_btn.pack(pady=20)
getdatareport(1)
# Run the Tkinter event loop
root.mainloop()
